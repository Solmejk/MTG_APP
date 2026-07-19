from docx import Document
from docx.shared import Mm, Pt
from pathlib import Path
from PIL import Image
from io import BytesIO
import json

class ProxyPrinter:
    def __init__(self):
        self.settings_file = "tokenprinter_settings.json"
        self.load_settings()
        self.input_folder = self.settings.get("default_input", "")
        self.output_folder = self.settings.get("default_output", "")
    
    def load_settings(self):
        try:
            if Path(self.settings_file).exists():
                with open(self.settings_file, 'r') as f:
                    self.settings = json.load(f)
            else:
                self.settings = {
                    "default_input": "",
                    "default_output": ""
                }
        except Exception as e:
            self.settings = {
                "default_input": "",
                "default_output": ""
            }
    
    def save_settings(self):
        try:
            with open(self.settings_file, 'w') as f:
                json.dump(self.settings, f, indent=4)
        except Exception as e:
            print("ERROR: Could not save settings")
    
    def load_images():
        # TODO
        images = []
        return images
    
    def create_printable_document(self):
        filename = "TODO"
        self.input_folder = self.input_entry.get()
        self.output_folder = self.output_entry.get()
    
        images = self.load_images()
        
        doc = Document()
        
        section = doc.sections[0]
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)

        images_per_row = 2
        
        for i in range(0, len(images), images_per_row):
            row_images = images[i:i + images_per_row]
            
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1
            
            for image in row_images:
                img = Image.open(image)
                rotated_img = img.rotate(-90, expand=True)
                
                if rotated_img.mode == 'RGBA':
                    rgb_img = Image.new('RGB', rotated_img.size, (255, 255, 255))
                    rgb_img.paste(rotated_img, mask=rotated_img.split()[3])
                    rotated_img = rgb_img
                
                img_buffer = BytesIO()
                rotated_img.save(img_buffer, format='JPEG')
                img_buffer.seek(0)
                
                run = p.add_run()
                run.add_picture(img_buffer, width=Mm(88))
        
        output_path = Path(self.output_folder) / f"{filename}.docx"
        doc.save(str(output_path))
            
if __name__ == "__main__":
    app = ProxyPrinter()